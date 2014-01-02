# encoding: utf-8

"""
Test suite for the docx.api module
"""

import pytest

from docx.api import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.package import Package
from docx.parts.document import DocumentPart, InlineShapes

from .unitutil import class_mock, instance_mock, method_mock, var_mock


class DescribeDocument(object):

    def it_opens_a_docx_on_construction(self, init_fixture):
        docx_, open_ = init_fixture
        document = Document(docx_)
        open_.assert_called_once_with(docx_)
        assert isinstance(document, Document)

    def it_can_open_a_docx_file(self, open_fixture):
        docx_, Package_, package_, document_part_ = open_fixture
        document_part, package = Document._open(docx_)
        Package_.open.assert_called_once_with(docx_)
        assert document_part is document_part
        assert package is package_

    def it_opens_default_template_if_no_file_provided(
            self, Package_, default_docx_):
        Document._open(None)
        Package_.open.assert_called_once_with(default_docx_)

    def it_should_raise_if_not_a_Word_file(self, Package_, package_, docx_):
        package_.main_document.content_type = 'foobar'
        with pytest.raises(ValueError):
            Document._open(docx_)

    def it_provides_access_to_the_document_body(self, document):
        body = document.body
        assert body is document._document_part.body

    def it_provides_access_to_the_document_inline_shapes(self, document):
        body = document.inline_shapes
        assert body is document._document_part.inline_shapes

    def it_can_add_an_inline_picture(self, add_picture_fixture):
        document, inline_shapes, image_path_or_stream_, inline_picture_ = (
            add_picture_fixture
        )
        inline_picture = document.add_inline_picture(image_path_or_stream_)
        inline_shapes.add_picture.assert_called_once_with(
            image_path_or_stream_
        )
        assert inline_picture is inline_picture_

    def it_can_save_the_package(self, save_fixture):
        document, package_, file_ = save_fixture
        document.save(file_)
        package_.save.assert_called_once_with(file_)

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def add_picture_fixture(self, request, open_, document_part_):
        document = Document()
        inline_shapes = instance_mock(request, InlineShapes)
        document_part_.inline_shapes = inline_shapes
        image_path_ = instance_mock(request, str)
        picture_shape_ = inline_shapes.add_picture.return_value
        return document, inline_shapes, image_path_, picture_shape_

    @pytest.fixture
    def default_docx_(self, request):
        return var_mock(request, 'docx.api._default_docx_path')

    @pytest.fixture
    def document(self, open_):
        return Document()

    @pytest.fixture
    def document_part_(self, request):
        return instance_mock(
            request, DocumentPart, content_type=CT.WML_DOCUMENT_MAIN
        )

    @pytest.fixture
    def docx_(self, request):
        return instance_mock(request, str)

    @pytest.fixture
    def init_fixture(self, docx_, open_):
        return docx_, open_

    @pytest.fixture
    def open_(self, request, document_part_, package_):
        return method_mock(
            request, Document, '_open',
            return_value=(document_part_, package_)
        )

    @pytest.fixture
    def open_fixture(self, docx_, Package_, package_, document_part_):
        return docx_, Package_, package_, document_part_

    @pytest.fixture
    def Package_(self, request, package_):
        Package_ = class_mock(request, 'docx.api.Package')
        Package_.open.return_value = package_
        return Package_

    @pytest.fixture
    def package_(self, request, document_part_):
        package_ = instance_mock(request, Package)
        package_.main_document = document_part_
        return package_

    @pytest.fixture
    def save_fixture(self, request, open_, package_):
        file_ = instance_mock(request, str)
        document = Document()
        return document, package_, file_
